"""
Text extraction for the two formats this service accepts.

Kept in its own module (not folded into services/document_service.py)
because extraction is purely a content-transformation concern -- it
doesn't know about HTTP, storage, or even our metadata shape. If a
third format gets added later (say .txt or .md), it lives here.
"""
from io import BytesIO

from pypdf import PdfReader
from docx import Document


class ExtractionError(Exception):
    """Raised when a file can't be parsed -- corrupt PDF, password-protected
    DOCX, etc. The router turns this into a 422 instead of a 500."""


def extract_pdf(raw: bytes) -> str:
    """Extract text from PDF bytes. Skips pages that yield None
    (image-only / unparseable pages) rather than failing the whole doc."""
    try:
        reader = PdfReader(BytesIO(raw))
    except Exception as e:
        raise ExtractionError(f"Could not read PDF: {e}") from e

    pages = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception:
            text = ""
        if text.strip():
            pages.append(text)
    return "\n\n".join(pages)


def extract_docx(raw: bytes) -> str:
    """Extract text from a .docx by joining all paragraph runs.
    Tables are flattened into pipe-separated lines so their cell text
    is still recoverable in the output."""
    try:
        doc = Document(BytesIO(raw))
    except Exception as e:
        raise ExtractionError(f"Could not read DOCX: {e}") from e

    parts = []
    for paragraph in doc.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)
    return "\n".join(parts)


def extract(raw: bytes, kind: str) -> str:
    if kind == "pdf":
        return extract_pdf(raw)
    if kind == "docx":
        return extract_docx(raw)
    raise ExtractionError(f"Unsupported kind: {kind}")
